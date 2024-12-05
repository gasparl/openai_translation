from openai import OpenAI
from openai import OpenAIError
from docx import Document
from docx.shared import Pt
import os
import logging
import time
import tiktoken
import json

# Set the model name
MODEL_NAME = 'gpt-4'

# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Path to the config.json file
config_path = os.path.join(script_dir, 'config.json')

# Load the configuration file
try:
    with open(config_path, 'r') as config_file:
        config = json.load(config_file)
except FileNotFoundError:
    raise FileNotFoundError(f"Configuration file not found at {config_path}. Please ensure 'config.json' exists.")

# Initialize the OpenAI client with your API key
client = OpenAI(api_key=config.get('OPENAI_API_KEY'))

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def read_docx(file_path):
    """
    Reads a .docx file and returns the text as a list of paragraphs.
    """
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs]
    return paragraphs

def estimate_tokens(text, model=MODEL_NAME):
    """
    Estimates the number of tokens in a text using tiktoken.
    """
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

def split_text(text_list, max_tokens=2048, model=MODEL_NAME):
    """
    Splits text into chunks that fit within the token limit.
    """
    chunks = []
    current_chunk = ''
    for paragraph in text_list:
        if paragraph.strip() == '':
            continue  # Skip empty paragraphs
        # Estimate tokens if we add this paragraph
        new_chunk = current_chunk + paragraph + '\n'
        if estimate_tokens(new_chunk, model) > max_tokens:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = paragraph + '\n'
        else:
            current_chunk = new_chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    return chunks

def translate_text(text, previous_translations='', previous_texts='', sample_translation='', source_lang='English', target_lang='Hungarian', model=MODEL_NAME):
    """
    Translates text using the OpenAI API, including previous translations and original texts.
    """
    # Build the system prompt with detailed instructions
    system_prompt = (
        f"You are a professional translator proficient in {source_lang} and {target_lang}."
        f" Your task is to translate the text with an emphasis on formal equivalence."
        f" Please preserve the original meaning, style, and sentence structure as closely as possible."
    )
    if sample_translation:
        system_prompt += f"\n\nUse the following sample translation as a style guide:\n\n{sample_translation}\n\n"

    # Build the user prompt
    prompt = f"Translate the following text from {source_lang} to {target_lang} with emphasis on formal equivalence."

    if previous_texts and previous_translations:
        prompt += "\n\nPrevious segments for context (original and translation):\n"
        # Combine previous texts and translations as pairs
        previous_texts_lines = previous_texts.strip().split('\n')
        previous_translations_lines = previous_translations.strip().split('\n')
        previous_pairs = zip(previous_texts_lines, previous_translations_lines)
        for orig, trans in previous_pairs:
            prompt += f"\nOriginal: {orig}\nTranslation: {trans}\n"

    prompt += f"\n\nText to translate:\n\n{text}"

    # Ensure total tokens are within limit
    max_model_tokens = 8192 if '32k' not in model else 32768
    max_completion_tokens = 1024  # Reserve tokens for the completion
    max_prompt_tokens = max_model_tokens - max_completion_tokens

    # Estimate tokens
    total_prompt = system_prompt + '\n\n' + prompt
    total_tokens = estimate_tokens(total_prompt, model)

    # Truncate previous translations and texts if necessary
    encoding = tiktoken.encoding_for_model(model)
    while total_tokens > max_prompt_tokens:
        if previous_texts and previous_translations:
            # Remove the oldest pair
            previous_texts_lines = previous_texts.strip().split('\n')
            previous_translations_lines = previous_translations.strip().split('\n')
            if len(previous_texts_lines) > 1 and len(previous_translations_lines) > 1:
                previous_texts = '\n'.join(previous_texts_lines[1:])
                previous_translations = '\n'.join(previous_translations_lines[1:])
            else:
                # Can't reduce previous context further
                previous_texts = ''
                previous_translations = ''
        else:
            # Need to truncate the text
            logging.warning("Prompt is too long even without previous translations. Truncating text.")
            text_token_ids = encoding.encode(text)
            allowed_text_tokens = max_prompt_tokens - estimate_tokens(system_prompt + '\n\n' + prompt.replace(text, ''), model)
            if allowed_text_tokens > 0:
                # Truncate text
                text = encoding.decode(text_token_ids[:allowed_text_tokens])
                # Rebuild prompt
                prompt = f"Translate the following text from {source_lang} to {target_lang} with emphasis on formal equivalence."
                if previous_texts and previous_translations:
                    prompt += "\n\nPrevious segments for context (original and translation):\n"
                    previous_pairs = zip(previous_texts.strip().split('\n'), previous_translations.strip().split('\n'))
                    for orig, trans in previous_pairs:
                        prompt += f"\nOriginal: {orig}\nTranslation: {trans}\n"
                prompt += f"\n\nText to translate:\n\n{text}"
                total_prompt = system_prompt + '\n\n' + prompt
                total_tokens = estimate_tokens(total_prompt, model)
            else:
                raise ValueError("Text to translate is too long to fit into the prompt.")
        # Recalculate total tokens
        total_prompt = system_prompt + '\n\n' + prompt
        total_tokens = estimate_tokens(total_prompt, model)

    # API call with error handling and retries
    max_retries = 5
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_completion_tokens,
                temperature=0.3,
            )
            translated_text = response.choices[0].message.content
            return translated_text.strip()
        except OpenAIError as e:
            logging.error(f"OpenAI API error: {e}")
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logging.info(f"Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                raise
        except Exception as e:
            logging.error(f"An unexpected error occurred: {e}")
            raise

def write_docx(paragraphs, output_path):
    """
    Writes a list of paragraphs to a .docx file.
    """
    doc = Document()
    for para in paragraphs:
        if para.strip() == '':
            doc.add_paragraph()
        else:
            doc_para = doc.add_paragraph(para)
            doc_para.style.font.name = 'Calibri'
            doc_para.style.font.size = Pt(12)
    doc.save(output_path)

def main(input_file, output_file, sample_translation_file=None):
    # Step 1: Read the original document
    logging.info(f"Reading input file: {input_file}")
    original_paragraphs = read_docx(input_file)

    # Step 2: Read the sample translation if provided
    sample_translation = ''
    if sample_translation_file and os.path.exists(sample_translation_file):
        logging.info(f"Reading sample translation file: {sample_translation_file}")
        sample_paragraphs = read_docx(sample_translation_file)
        sample_translation = '\n'.join(sample_paragraphs)

    # Step 3: Split the text into chunks
    logging.info("Splitting text into chunks...")
    # Adjust max_tokens for chunks, considering the token limits and the amount of context
    chunk_max_tokens = 2048  # You can adjust this value
    text_chunks = split_text(original_paragraphs, max_tokens=chunk_max_tokens)

    # Step 4: Translate each chunk with context
    translated_paragraphs = []
    previous_translations = ''
    previous_texts = ''
    max_previous_segments = 5  # Adjust based on your preference
    for idx, chunk in enumerate(text_chunks):
        logging.info(f"Translating chunk {idx+1}/{len(text_chunks)}...")
        try:
            translated_chunk = translate_text(
                text=chunk,
                previous_translations=previous_translations,
                previous_texts=previous_texts,
                sample_translation=sample_translation
            )
        except Exception as e:
            logging.error(f"Failed to translate chunk {idx+1}: {e}")
            continue
        # Update previous texts and translations
        previous_texts += '\n' + chunk
        previous_translations += '\n' + translated_chunk
        # Keep only the last few segments
        previous_texts_lines = previous_texts.strip().split('\n')
        previous_translations_lines = previous_translations.strip().split('\n')
        if len(previous_texts_lines) > max_previous_segments:
            previous_texts = '\n'.join(previous_texts_lines[-max_previous_segments:])
            previous_translations = '\n'.join(previous_translations_lines[-max_previous_segments:])
        # Add the translated paragraphs to the list
        translated_chunk_paragraphs = translated_chunk.strip().split('\n')
        translated_paragraphs.extend(translated_chunk_paragraphs)
    # Step 5: Write the translated text to a new document
    logging.info(f"Writing translated text to output file: {output_file}")
    write_docx(translated_paragraphs, output_file)
    logging.info(f"Translation completed. The translated document is saved as '{output_file}'.")

if __name__ == "__main__":
    # Replace 'input.docx' with your input file path,
    # 'output.docx' with your desired output file path,
    # and 'sample_translation.docx' with the path to your sample translation (optional)
    # main('input.docx', 'output.docx', sample_translation_file='sample_translation.docx')
    main('input.docx', 'output.docx')
